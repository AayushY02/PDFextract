"""
Deterministic Layout + Tables Document Text Extractor (v2)
=========================================================

Keeps the SAME output contract as your current extractor:
- Recursively scans INPUT_DIR for PDF, Word (.docx), Excel (.xls/.xlsx)
- Extracts embedded text (NO OCR) and saves to results/output1/ preserving subfolders
- Adds page boundary identifiers ONLY:
    [[PAGE_START N]] at the beginning of each page (then a blank line)
    [[PAGE_END]] at the end of each page
- Detects tabular regions (PDF/DOCX/Excel) and replaces their text with structured markers:
    [[TABLE_START page=P index=I rows=R cols=C]]
    col1 | col2 | col3
    [[TABLE_END]]
- Heading marker:
    [[HEADING]] <line>
  (based on stable font-size threshold per page)

Big improvements vs PyMuPDF block-order extraction:
- Uses pdfplumber/pdfminer layout geometry (chars with coordinates)
- Builds lines deterministically (no mixed extraction modes)
- Finds tables with bbox support and removes their chars from normal text flow
- Outputs tables in correct reading order position

Requirements:
    pip install pdfplumber pdfminer.six python-docx openpyxl pandas

Usage:
    python text_extractor_layout_tables_v2.py
"""

from __future__ import annotations

import csv
import math
import statistics
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
from concurrent.futures import ProcessPoolExecutor, as_completed
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import pdfplumber

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# ---------- CONFIGURATION ----------
INPUT_DIR = Path("input")
RESULTS_DIR = Path("results")
OUTPUT_DIR = RESULTS_DIR / "output12"
LOG_FILE = RESULTS_DIR / "extraction_log.csv"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls"}

PAGE_START_TEMPLATE = "[[PAGE_START {n}]]"
PAGE_END_TEMPLATE = "[[PAGE_END]]"

TABLE_COL_SEPARATOR = " | "

# ---- PDF line building knobs (tune if needed) ----
LINE_Y_TOLERANCE = 3.0          # pts: how tightly to cluster chars into a line by y/top
SPACE_GAP_FACTOR = 0.6          # heuristic for inserting spaces in Latin text based on char size
TABLE_BBOX_PADDING = 1.5        # pts: expand detected table bboxes for safer suppression
MIN_TABLE_ROWS = 2
MIN_TABLE_COLS = 2

# ---- PDF table detection settings ----
# LATTICE (works when there are ruling lines)
TABLE_SETTINGS_LATTICE = {
    "vertical_strategy": "lines",
    "horizontal_strategy": "lines",
    "intersection_tolerance": 5,
    "snap_tolerance": 3,
    "join_tolerance": 3,
    "edge_min_length": 3,
    "min_words_vertical": 1,
    "min_words_horizontal": 1,
    "text_tolerance": 3,
}

# STREAM (works when it's aligned text without ruling lines)
TABLE_SETTINGS_STREAM = {
    "vertical_strategy": "text",
    "horizontal_strategy": "text",
    "intersection_tolerance": 3,
    "snap_tolerance": 3,
    "join_tolerance": 3,
    "edge_min_length": 3,
    "min_words_vertical": 2,
    "min_words_horizontal": 1,
    "text_tolerance": 3,
}
# ----------------------------------


# ---------- Helpers for formatting ----------
def sanitize_cell_text(cell_value: Any) -> str:
    """Normalizes table cell text for compact export."""
    if cell_value is None:
        return ""
    s = str(cell_value)
    if not s.strip():
        return ""
    return " ".join(s.split())


def format_table_rows(rows: List[List[Any]], header_label: str) -> str:
    """Convert row data into a structured table block with explicit markers."""
    clean_rows = [[sanitize_cell_text(c) for c in row] for row in rows]
    row_count = len(clean_rows)
    col_count = max((len(r) for r in clean_rows), default=0)

    header = f"[[TABLE_START {header_label} rows={row_count} cols={col_count}]]"
    body_lines = []
    for r in clean_rows:
        # pad shorter rows so delimiters stay consistent
        if len(r) < col_count:
            r = r + [""] * (col_count - len(r))
        body_lines.append(TABLE_COL_SEPARATOR.join(r))
    footer = "[[TABLE_END]]"

    out = [header]
    if body_lines:
        out.extend(body_lines)
    out.append(footer)
    return "\n".join(out) + "\n\n"


# ---------- DOCX handling (same rules as your current script) ----------
def iter_docx_blocks(document: Document):
    """Yield paragraphs and tables in original document order."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def extract_text_from_docx(docx_path: Path) -> Tuple[str, int]:
    doc = Document(docx_path)
    segments: List[str] = []
    table_index = 0

    for block in iter_docx_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            line = text + "\n"
            style_name = block.style.name if block.style else ""
            if style_name and style_name.lower().startswith("heading"):
                segments.append(f"[[HEADING]] {line}")
            else:
                segments.append(line)

        elif isinstance(block, Table):
            table_index += 1
            rows = []
            for row in block.rows:
                rows.append([cell.text for cell in row.cells])
            segments.append(format_table_rows(rows, header_label=f"page=1 index={table_index}"))

    start_marker = f"{PAGE_START_TEMPLATE.format(n=1)}\n\n"
    end_marker = f"\n{PAGE_END_TEMPLATE}\n"
    return start_marker + "".join(segments) + end_marker, 1


# ---------- Excel handling (same rules as your current script) ----------
def extract_text_from_excel(excel_path: Path) -> Tuple[str, int]:
    workbook = pd.ExcelFile(excel_path)
    chunks: List[str] = []

    for idx, sheet_name in enumerate(workbook.sheet_names, start=1):
        df = workbook.parse(sheet_name, header=None, dtype=object).where(pd.notnull, "")
        rows = df.astype(str).values.tolist()

        table_text = format_table_rows(
            rows=rows,
            header_label=f"page={idx} index=1 sheet={sanitize_cell_text(sheet_name)}",
        )

        start_marker = f"{PAGE_START_TEMPLATE.format(n=idx)}\n\n"
        end_marker = f"\n{PAGE_END_TEMPLATE}\n"
        chunks.append(start_marker + table_text + end_marker)

    return "".join(chunks), len(workbook.sheet_names)


# ---------- PDF layout + tables ----------
@dataclass
class PDFLine:
    top: float
    x0: float
    x1: float
    bottom: float
    text: str
    max_size: float


@dataclass
class PDFTable:
    top: float
    x0: float
    bbox: Tuple[float, float, float, float]  # (x0, top, x1, bottom)
    content: str


def _pad_bbox(b: Tuple[float, float, float, float], pad: float) -> Tuple[float, float, float, float]:
    x0, top, x1, bottom = b
    return (x0 - pad, top - pad, x1 + pad, bottom + pad)


def _point_in_bbox(x: float, y: float, b: Tuple[float, float, float, float]) -> bool:
    x0, top, x1, bottom = b
    return (x0 <= x <= x1) and (top <= y <= bottom)


def _is_ascii_word_char(ch: str) -> bool:
    # helps decide spacing for Latin text; Japanese should usually not get inserted spaces
    return ch.isascii() and (ch.isalnum() or ch in {"_", "-", "."})


def _build_lines_from_chars(chars: List[Dict[str, Any]]) -> List[PDFLine]:
    """
    Turn pdfplumber chars into deterministic lines:
    - cluster by 'top' (within LINE_Y_TOLERANCE)
    - sort within line by x0
    - join chars, with cautious space insertion for Latin words only
    """
    if not chars:
        return []

    # Sort for stable clustering
    chars_sorted = sorted(chars, key=lambda c: (c.get("top", 0.0), c.get("x0", 0.0)))

    lines_chars: List[List[Dict[str, Any]]] = []
    current: List[Dict[str, Any]] = []
    current_top: Optional[float] = None

    for ch in chars_sorted:
        top = float(ch.get("top", 0.0))
        if current_top is None:
            current = [ch]
            current_top = top
            continue

        if abs(top - current_top) <= LINE_Y_TOLERANCE:
            current.append(ch)
            # keep a running average so we don't drift line boundaries
            current_top = (current_top * (len(current) - 1) + top) / len(current)
        else:
            lines_chars.append(current)
            current = [ch]
            current_top = top

    if current:
        lines_chars.append(current)

    lines: List[PDFLine] = []
    for group in lines_chars:
        group_sorted = sorted(group, key=lambda c: float(c.get("x0", 0.0)))

        # compute dynamic space threshold
        sizes = [float(c.get("size", 0.0)) for c in group_sorted if isinstance(c.get("size", None), (int, float))]
        base_size = statistics.median(sizes) if sizes else 10.0
        space_threshold = base_size * SPACE_GAP_FACTOR

        parts: List[str] = []
        max_size = 0.0
        x0 = math.inf
        x1 = -math.inf
        top = math.inf
        bottom = -math.inf

        prev = None
        for c in group_sorted:
            t = str(c.get("text", ""))
            if not t:
                continue

            cx0 = float(c.get("x0", 0.0))
            cx1 = float(c.get("x1", 0.0))
            ctop = float(c.get("top", 0.0))
            cbottom = float(c.get("bottom", 0.0))
            csize = float(c.get("size", 0.0)) if isinstance(c.get("size", None), (int, float)) else 0.0
            max_size = max(max_size, csize)

            x0 = min(x0, cx0)
            x1 = max(x1, cx1)
            top = min(top, ctop)
            bottom = max(bottom, cbottom)

            if prev is not None:
                gap = cx0 - float(prev.get("x1", cx0))
                # insert space ONLY for Latin-ish sequences; don't mess with Japanese
                if gap > space_threshold:
                    prev_t = str(prev.get("text", ""))[-1:] if prev.get("text") else ""
                    if _is_ascii_word_char(prev_t) and _is_ascii_word_char(t[:1]):
                        if parts and not parts[-1].endswith(" "):
                            parts.append(" ")

            parts.append(t)
            prev = c

        text = "".join(parts)
        if text.strip():
            if not text.endswith("\n"):
                text += "\n"
            lines.append(PDFLine(top=top, x0=x0, x1=x1, bottom=bottom, text=text, max_size=max_size))

    return lines


def _extract_tables_pdf(page: pdfplumber.page.Page, page_number: int) -> List[PDFTable]:
    """
    Find tables with bbox support.
    Uses two passes: lattice (lines) then stream (text) and merges results.
    """
    candidates = []

    def add_tables(table_settings: Dict[str, Any]):
        try:
            found = page.find_tables(table_settings=table_settings)  # list of Table objects
        except Exception:
            return
        for i, t in enumerate(found, start=1):
            rows = t.extract()
            if not rows:
                continue
            # quick sanity filter
            rcount = len(rows)
            ccount = max((len(r) for r in rows), default=0)
            if rcount < MIN_TABLE_ROWS or ccount < MIN_TABLE_COLS:
                continue

            # density score helps choose better duplicate
            nonempty = sum(1 for r in rows for c in r if sanitize_cell_text(c))
            total = max(rcount * ccount, 1)
            density = nonempty / total

            bbox = tuple(t.bbox)  # (x0, top, x1, bottom) in pdfplumber coords
            candidates.append(
                {
                    "bbox": bbox,
                    "rows": rows,
                    "rcount": rcount,
                    "ccount": ccount,
                    "density": density,
                }
            )

    add_tables(TABLE_SETTINGS_LATTICE)
    add_tables(TABLE_SETTINGS_STREAM)

    # Deduplicate overlapping candidates by IoU and keep best by (cells, density)
    def iou(a, b) -> float:
        ax0, at, ax1, ab = a
        bx0, bt, bx1, bb = b
        ix0, it = max(ax0, bx0), max(at, bt)
        ix1, ib = min(ax1, bx1), min(ab, bb)
        iw, ih = max(0.0, ix1 - ix0), max(0.0, ib - it)
        inter = iw * ih
        area_a = max(0.0, (ax1 - ax0)) * max(0.0, (ab - at))
        area_b = max(0.0, (bx1 - bx0)) * max(0.0, (bb - bt))
        union = max(area_a + area_b - inter, 1e-9)
        return inter / union

    selected = []
    for cand in sorted(candidates, key=lambda c: (c["rcount"] * c["ccount"], c["density"]), reverse=True):
        keep = True
        for s in selected:
            if iou(cand["bbox"], s["bbox"]) > 0.85:
                keep = False
                break
        if keep:
            selected.append(cand)

    tables_out: List[PDFTable] = []
    for idx, t in enumerate(sorted(selected, key=lambda c: (c["bbox"][1], c["bbox"][0])), start=1):
        bbox_padded = _pad_bbox(t["bbox"], TABLE_BBOX_PADDING)
        header_label = f"page={page_number} index={idx}"
        content = format_table_rows(t["rows"], header_label=header_label)
        x0, top, _, _ = bbox_padded
        tables_out.append(PDFTable(top=top, x0=x0, bbox=bbox_padded, content=content))

    return tables_out


def _extract_page_body_pdf(page: pdfplumber.page.Page, page_number: int) -> str:
    # 1) Find tables (with bbox)
    tables = _extract_tables_pdf(page, page_number)

    # 2) Get chars and suppress those inside table bboxes
    chars = page.chars or []
    if tables:
        filtered = []
        for c in chars:
            cx = (float(c.get("x0", 0.0)) + float(c.get("x1", 0.0))) / 2.0
            cy = (float(c.get("top", 0.0)) + float(c.get("bottom", 0.0))) / 2.0
            if any(_point_in_bbox(cx, cy, t.bbox) for t in tables):
                continue
            filtered.append(c)
        chars = filtered

    # 3) Build deterministic lines from remaining chars
    lines = _build_lines_from_chars(chars)

    # 4) Heading threshold (stable)
    sizes = [ln.max_size for ln in lines if ln.max_size > 0]
    threshold = (statistics.median(sizes) + 1.5) if sizes else float("inf")

    # 5) Merge lines + tables into one reading-order stream
    items: List[Tuple[float, float, str, str]] = []
    # tuple: (top, x0, kind, content)
    for t in tables:
        items.append((t.top, t.x0, "table", t.content))
    for ln in lines:
        txt = ln.text
        collapsed = "".join(txt.strip().split())
        if collapsed and ln.max_size >= threshold:
            items.append((ln.top, ln.x0, "line", f"[[HEADING]] {txt}"))
        else:
            items.append((ln.top, ln.x0, "line", txt))

    # Optional: detect two-column layout and order col1 then col2
    # We use line x0 distribution only (excluding tables) to avoid skew.
    line_xs = sorted(x for (top, x, kind, _) in items if kind == "line")
    split_x = None
    if len(line_xs) >= 30:
        # find largest gap in x0 positions
        gaps = [(line_xs[i + 1] - line_xs[i], i) for i in range(len(line_xs) - 1)]
        best_gap, idx = max(gaps, key=lambda g: g[0])
        # conservative thresholds
        if best_gap > 90 and 0.2 * len(line_xs) < idx < 0.8 * len(line_xs):
            split_x = (line_xs[idx] + line_xs[idx + 1]) / 2.0

    if split_x is None:
        items.sort(key=lambda it: (it[0], it[1]))
    else:
        # left column then right column
        items.sort(key=lambda it: ((0 if it[1] < split_x else 1), it[0], it[1]))

    page_text = "".join(content for (_, _, _, content) in items)
    if page_text and not page_text.endswith("\n"):
        page_text += "\n"
    return page_text


def extract_text_from_pdf(pdf_path: Path) -> Tuple[str, int]:
    with pdfplumber.open(str(pdf_path)) as pdf:
        chunks: List[str] = []
        page_count = len(pdf.pages)

        for idx, page in enumerate(pdf.pages, start=1):
            page_body = _extract_page_body_pdf(page, idx)

            start_marker = f"{PAGE_START_TEMPLATE.format(n=idx)}\n\n"
            end_marker = f"\n{PAGE_END_TEMPLATE}\n"
            chunks.append(start_marker + (page_body or "") + end_marker)

        return "".join(chunks), page_count


# ---------- IO / Logging / Driver ----------
def save_text(output_path: Path, text: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


def init_log() -> None:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    with open(LOG_FILE, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["timestamp", "file", "pages", "status", "error"])


def append_log(file_path: Path, pages: int, status: str, error: str = "") -> None:
    with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            datetime.now().isoformat(timespec="seconds"),
            str(file_path),
            pages,
            status,
            error
        ])


def collect_supported_files(input_dir: Path) -> List[Path]:
    return sorted(
        (
            p for p in input_dir.rglob("*")
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
        ),
        key=lambda p: str(p).lower(),
    )


def process_one_file(file_path: Path, input_dir: Path, output_dir: Path):
    rel = file_path.relative_to(input_dir)
    out_path = output_dir / rel.with_suffix(".txt")
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            text, pages = extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            text, pages = extract_text_from_docx(file_path)
        elif suffix in {".xlsx", ".xls"}:
            text, pages = extract_text_from_excel(file_path)
        else:
            raise ValueError(f"Unsupported extension: {suffix}")

        save_text(out_path, text)
        return rel, pages, "OK", ""

    except Exception as e:
        return rel, 0, "FAIL", str(e)


def process_documents(input_dir: Path, output_dir: Path) -> None:
    files = collect_supported_files(input_dir)
    if not files:
        print(f"No supported files found in: {input_dir}")
        return

    init_log()
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Found {len(files)} files. Starting extraction...\n")

    with ProcessPoolExecutor() as executor:
        futures = {
            executor.submit(process_one_file, p, input_dir, output_dir): p
            for p in files
        }

        for fut in as_completed(futures):
            rel, pages, status, error = fut.result()
            append_log(rel, pages, status, error)
            if status == "OK":
                print(f"✅ Extracted: {rel}  ({pages} pages/sheets)")
            else:
                print(f"❌ Failed: {rel} ({error})")

    print("\nAll done!")
    print(f"📂 Extracted texts: {output_dir}")
    print(f"📄 Log file: {LOG_FILE}")


# Backwards-compatible entry point name
def process_pdfs(input_dir: Path, output_dir: Path) -> None:
    process_documents(input_dir, output_dir)


if __name__ == "__main__":
    process_documents(INPUT_DIR, OUTPUT_DIR)
