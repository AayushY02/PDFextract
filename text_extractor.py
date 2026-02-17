"""
Document Text Extractor with Page Start/End Markers (results/output1 version)
=============================================================================

This script recursively scans the 'input' folder for PDF, Word (.docx),
and Excel (.xls/.xlsx) files, extracts their embedded text (no OCR),
and saves the result in:
    results/output1/

Enhancements:
    - Adds explicit page boundary identifiers ONLY:
        * [[PAGE_START N]] at the very beginning of each page (then a blank line)
        * [[PAGE_END   N]] after each page's text
    - No "Page N" footer lines
    - No form-feed (\f) separators
    - Detects tabular regions and replaces their text with structured markers:
        * [[TABLE_START page=P index=I rows=R cols=C]]
        * Pipe-delimited rows (`col1 | col2 | col3`)
        * [[TABLE_END]]

Layout per page N in the TXT:
    [[PAGE_START N]]

    <page N text>

    [[PAGE_END   N]]

It also writes a CSV log file (results/extraction_log.csv)
summarizing:
    - File path (relative to input)
    - Page count
    - Status (OK/FAIL)
    - Error message (if any)

Requirements:
    pip install PyMuPDF python-docx openpyxl pandas

Usage:
    python pdf_text_extractor.py
"""

import csv
import fitz  # PyMuPDF
import pandas as pd
import statistics
import re
import unicodedata
from pathlib import Path
from datetime import datetime
from concurrent.futures import ProcessPoolExecutor, as_completed
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
# ---------- CONFIGURATION ----------
INPUT_DIR = Path("input")
RESULTS_DIR = Path("results")
OUTPUT_DIR = RESULTS_DIR / "output1"
LOG_FILE = RESULTS_DIR / "extraction_log.csv"
QUALITY_LOG_FILE = RESULTS_DIR / "extraction_quality.csv"
QUALITY_LOG_FILE_EN = RESULTS_DIR / "extraction_quality_en.csv"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls"}

# Page boundary marker settings
PAGE_START_TEMPLATE = "[[PAGE_START {n}]]"
PAGE_END_TEMPLATE = "[[PAGE_END]]"
TABLE_COL_SEPARATOR = " | "

# ---------- QUALITY DETECTION CONFIG ----------
WEIRD_REPLACEMENT_THRESHOLD = 1
WEIRD_CONTROL_RATIO = 0.001
WEIRD_PRIVATE_RATIO = 0.002
WEIRD_NONEXPECTED_RATIO = 0.25

TABLE_RATIO_THRESHOLD = 0.70
TABLE_RATIO_SECONDARY = 0.50
TABLE_COUNT_THRESHOLD = 3

EMPTY_ENUM_RUN_MIN = 5

TABLE_START_TOKEN = "[[TABLE_START"
TABLE_END_TOKEN = "[[TABLE_END]]"
# ----------------------------------

def process_one_file(file_path: Path, input_dir: Path, output_dir: Path):
    relative_path = file_path.relative_to(input_dir)
    output_path = output_dir / relative_path.with_suffix(".txt")
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            text, pages = extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            text, pages = extract_text_from_docx(file_path)
        elif suffix in {".xlsx", ".xls"}:
            text, pages = extract_text_from_excel(file_path)
        else:
            raise ValueError(f"Unsupported extension: {suffix}")

        quality = analyze_text_quality(text, suffix)
        save_text(output_path, text)
        return relative_path, pages, "OK", "", quality
    except Exception as e:
        return relative_path, 0, "FAIL", str(e), None


def sanitize_cell_text(cell_value):
    """
    Normalizes table cell text for compact export.
    Collapses internal whitespace and removes surrounding blanks.
    """
    if not cell_value:
        return ""
    return " ".join(str(cell_value).split())


def format_table_rows(rows, header_label: str, row_count: int, col_count: int) -> str:
    """
    Convert arbitrary row data into a structured text block with explicit markers.
    header_label allows callers to inject context such as page or sheet numbers.
    """
    header = f"[[TABLE_START {header_label} rows={row_count} cols={col_count}]]"
    body_rows = []
    for row in rows:
        sanitized_cells = [sanitize_cell_text(cell) for cell in row]
        body_rows.append(TABLE_COL_SEPARATOR.join(sanitized_cells))
    body = "\n".join(body_rows)
    footer = "[[TABLE_END]]"
    content_lines = [header]
    if body:
        content_lines.append(body)
    content_lines.append(footer)
    return "\n".join(content_lines) + "\n\n"


def format_table(table, page_number: int, table_index: int) -> str:
    """
    Convert a PyMuPDF table object into a structured text block with explicit markers.
    Cells are joined with the configured column separator to keep columns clear.
    """
    header_label = f"page={page_number} index={table_index}"
    return format_table_rows(
        table.extract(),
        header_label=header_label,
        row_count=table.row_count,
        col_count=table.col_count,
    )


def block_center_in_bbox(block_bbox, table_bbox, padding: float = 0.5) -> bool:
    """
    Check if the center of a text block falls within the table bounding box.
    Padding slightly enlarges the table bounds to cover rounding differences.
    """
    bx0, by0, bx1, by1 = block_bbox
    tx0, ty0, tx1, ty1 = table_bbox
    cx = (bx0 + bx1) / 2
    cy = (by0 + by1) / 2
    return (
        (tx0 - padding) <= cx <= (tx1 + padding)
        and (ty0 - padding) <= cy <= (ty1 + padding)
    )

def line_belongs_to_table(
    line_bbox,
    table_entry,
    padding: float = 0.5,
    bottom_margin_rows: float = 0.4,
) -> bool:
    """
    Decide if a text line should be treated as part of a table.

    We slightly shrink the bottom of the table bbox so that normal text just
    under the table is not accidentally considered "inside" the table region.
    """
    bx0, by0, bx1, by1 = line_bbox
    cx = (bx0 + bx1) / 2
    cy = (by0 + by1) / 2

    tx0, ty0, tx1, ty1 = table_entry["bbox"]
    row_h = table_entry.get("row_height") or 0.0

    # shrink table bottom by a fraction of one row height
    if row_h > 0:
        ty1_adj = ty1 - row_h * bottom_margin_rows
    else:
        ty1_adj = ty1

    return (
        (tx0 - padding) <= cx <= (tx1 + padding)
        and (ty0 - padding) <= cy <= (ty1_adj + padding)
    )


def iter_docx_blocks(document):
    """
    Yield paragraphs and tables in the original document order.
    """
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def extract_page_body(page, page_number: int) -> str:
    """
    Build the body text for a single page, replacing detected tables with structured markers.
    """
    table_finder = page.find_tables()
    tables = []
    # if table_finder:
    #     for table_index, table in enumerate(table_finder.tables, start=1):
    #         tables.append(
    #             {
    #                 "bbox": table.bbox,
    #                 "content": format_table(table, page_number, table_index),
    #                 "emitted": False,
    #             }
    #         )
    if table_finder:
        for table_index, table in enumerate(table_finder.tables, start=1):
            tx0, ty0, tx1, ty1 = table.bbox
            height = ty1 - ty0
            row_h = height / max(table.row_count, 1)

            tables.append(
                {
                    "bbox": table.bbox,
                    "row_height": row_h,
                    "content": format_table(table, page_number, table_index),
                    "emitted": False,
                }
            )


    segments = []
    info = page.get_text("dict")
    span_sizes = []
    for block in info.get("blocks", []):
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                sz = span.get("size")
                if isinstance(sz, (int, float)):
                    span_sizes.append(float(sz))
    median_size = statistics.median(span_sizes) if span_sizes else 0.0
    size_threshold = median_size + 1.5

    for block in info.get("blocks", []):
        for line in block.get("lines", []):
            # compose line text
            line_text = "".join(span.get("text", "") for span in line.get("spans", []))
            if not line_text:
                continue
            if not line_text.endswith("\n"):
                line_text += "\n"

            # line bbox and max size
            x0, y0, x1, y1 = line.get("bbox", [0, 0, 0, 0])
            line_bbox = (x0, y0, x1, y1)
            sizes = [float(span.get("size", 0)) for span in line.get("spans", [])]
            line_size = max(sizes) if sizes else 0.0

            # table handling: emit once when first encountered
            # table_entry = None
            # for entry in tables:
            #     if block_center_in_bbox(line_bbox, entry["bbox"]):
            #         table_entry = entry
            #         break
            # if table_entry is not None:
            #     if not table_entry["emitted"]:
            #         segments.append(table_entry["content"])
            #         table_entry["emitted"] = True
            #     continue
            table_entry = None
            for entry in tables:
                if line_belongs_to_table(line_bbox, entry):
                    table_entry = entry
                    break

            if table_entry is not None:
                # emit table once, skip individual table lines
                if not table_entry["emitted"]:
                    segments.append(table_entry["content"])
                    table_entry["emitted"] = True
                continue


            # heading marker for large font lines (inline before the text)
            collapsed = "".join(line_text.strip().split())
            if collapsed and line_size >= size_threshold:
                segments.append(f"[[HEADING]] {line_text}")
            else:
                segments.append(line_text)

    for entry in tables:
        if not entry["emitted"]:
            segments.append(entry["content"])

    page_body = "".join(segments)
    if page_body and not page_body.endswith("\n"):
        page_body += "\n"
    return page_body


def extract_text_from_pdf(pdf_path: Path):
    """
    Extracts text and counts pages from a PDF using PyMuPDF.

    For each page N:
        - Insert a page-start marker line: [[PAGE_START N]], followed by a blank line.
        - Append the page's extracted text (ensuring trailing newline).
        - Insert a page-end marker line: [[PAGE_END N]].

    Returns:
        combined_text (str), page_count (int)
    """
    with fitz.open(pdf_path) as doc:
        page_count = len(doc)
        chunks = []

        for idx, page in enumerate(doc, start=1):
            page_text = extract_page_body(page, idx)
            if not page_text:
                fallback_text = page.get_text("text") or ""
                if fallback_text and not fallback_text.endswith("\n"):
                    fallback_text += "\n"
                page_text = fallback_text

            # Build the page chunk with ONLY our markers
            start_marker = f"{PAGE_START_TEMPLATE.format(n=idx)}\n\n"
            end_marker = f"\n{PAGE_END_TEMPLATE.format(n=idx)}\n"

            chunk = start_marker + page_text + end_marker
            chunks.append(chunk)

        combined_text = "".join(chunks)
        return combined_text, page_count


def extract_text_from_docx(docx_path: Path):
    """
    Extract text from a Word document, preserving heading markers and tables.
    The whole document is treated as a single page for marker consistency.
    """
    doc = Document(docx_path)
    segments = []
    table_index = 0

    for block in iter_docx_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if not text.endswith("\n"):
                text = f"{text}\n"
            style_name = block.style.name if block.style else ""
            if style_name and style_name.lower().startswith("heading"):
                segments.append(f"[[HEADING]] {text}")
            else:
                segments.append(text)
        elif isinstance(block, Table):
            table_index += 1
            rows = []
            max_cols = 0
            for row in block.rows:
                cells = [sanitize_cell_text(cell.text) for cell in row.cells]
                max_cols = max(max_cols, len(cells))
                rows.append(cells)

            segments.append(
                format_table_rows(
                    rows=rows,
                    header_label=f"page=1 index={table_index}",
                    row_count=len(rows),
                    col_count=max_cols,
                )
            )

    body = "".join(segments)
    start_marker = f"{PAGE_START_TEMPLATE.format(n=1)}\n\n"
    end_marker = f"\n{PAGE_END_TEMPLATE.format(n=1)}\n"
    combined = start_marker + body + end_marker
    return combined, 1


def extract_text_from_excel(excel_path: Path):
    """
    Extract text from Excel workbooks by emitting one table per sheet.
    Each sheet is wrapped in page markers to align with the PDF output format.
    """
    workbook = pd.ExcelFile(excel_path)
    chunks = []

    for idx, sheet_name in enumerate(workbook.sheet_names, start=1):
        df = workbook.parse(sheet_name, header=None, dtype=object)
        df = df.where(pd.notnull(df), "")
        rows = df.astype(str).values.tolist()
        col_count = df.shape[1] if not df.empty else 0

        table_text = format_table_rows(
            rows=rows,
            header_label=f"page={idx} index=1 sheet={sanitize_cell_text(sheet_name)}",
            row_count=len(rows),
            col_count=col_count,
        )
        start_marker = f"{PAGE_START_TEMPLATE.format(n=idx)}\n\n"
        end_marker = f"\n{PAGE_END_TEMPLATE.format(n=idx)}\n"
        chunks.append(start_marker + table_text + end_marker)

    combined_text = "".join(chunks)
    return combined_text, len(workbook.sheet_names)


def save_text(output_path: Path, text: str):
    """Writes text to a UTF-8 encoded file."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


# ---------- QUALITY DETECTION ----------
def _strip_markers_for_quality(text: str) -> str:
    if "[[" not in text:
        return text
    text = re.sub(r"\[\[PAGE_START[^\]]*\]\]", "", text)
    text = re.sub(r"\[\[PAGE_END\]\]", "", text)
    text = re.sub(r"\[\[TABLE_START[^\]]*\]\]", "", text)
    text = text.replace("[[TABLE_END]]", "")
    text = text.replace("[[HEADING]]", "")
    return text


def _table_stats(text: str) -> dict:
    if TABLE_START_TOKEN not in text:
        total = len(text)
        return {
            "table_blocks": 0,
            "table_chars": 0,
            "non_table_chars": total,
            "table_ratio": 0.0,
        }

    i = 0
    table_chars = 0
    non_table_chars = 0
    table_blocks = 0
    text_len = len(text)

    while i < text_len:
        start = text.find(TABLE_START_TOKEN, i)
        if start == -1:
            non_table_chars += text_len - i
            break

        non_table_chars += max(0, start - i)
        table_blocks += 1

        # content starts after the marker line
        start_line_end = text.find("\n", start)
        content_start = start_line_end + 1 if start_line_end != -1 else start

        end = text.find(TABLE_END_TOKEN, start)
        if end == -1:
            table_chars += max(0, text_len - content_start)
            break

        table_chars += max(0, end - content_start)
        i = end + len(TABLE_END_TOKEN)

    total = table_chars + non_table_chars
    table_ratio = (table_chars / total) if total > 0 else 0.0
    return {
        "table_blocks": table_blocks,
        "table_chars": table_chars,
        "non_table_chars": non_table_chars,
        "table_ratio": table_ratio,
    }


def _remove_table_blocks(text: str) -> str:
    if TABLE_START_TOKEN not in text:
        return text
    out = []
    i = 0
    n = len(text)
    while i < n:
        start = text.find(TABLE_START_TOKEN, i)
        if start == -1:
            out.append(text[i:])
            break
        out.append(text[i:start])
        end = text.find(TABLE_END_TOKEN, start)
        if end == -1:
            break
        i = end + len(TABLE_END_TOKEN)
    return "".join(out)


def _max_empty_circled_enum_run(text: str) -> int:
    circled = {chr(code) for code in range(0x2460, 0x2474)}  # ①..⑳
    cleaned = _strip_markers_for_quality(_remove_table_blocks(text))
    max_run = 0
    current = 0
    for line in cleaned.splitlines():
        s = line.strip()
        if s in circled:
            current += 1
            if current > max_run:
                max_run = current
        else:
            current = 0
    return max_run


def _is_expected_char(ch: str) -> bool:
    if ch in ("\n", "\r", "\t"):
        return True
    code = ord(ch)
    if 0x20 <= code <= 0x7E:
        return True  # ASCII printable
    if 0x3000 <= code <= 0x303F:
        return True  # CJK punctuation
    if 0x3040 <= code <= 0x309F:
        return True  # Hiragana
    if 0x30A0 <= code <= 0x30FF:
        return True  # Katakana
    if 0xFF66 <= code <= 0xFF9D:
        return True  # Half-width Katakana
    if 0x4E00 <= code <= 0x9FFF:
        return True  # CJK Unified
    if 0x3400 <= code <= 0x4DBF:
        return True  # CJK Ext A
    if 0xF900 <= code <= 0xFAFF:
        return True  # CJK Compatibility
    if 0xFF01 <= code <= 0xFF60:
        return True  # Full-width ASCII variants
    return False


def analyze_text_quality(text: str, source_ext: str) -> dict:
    cleaned = _strip_markers_for_quality(text)
    total_chars = len(cleaned)

    replacement_count = 0
    control_count = 0
    private_use_count = 0
    expected_count = 0

    for ch in cleaned:
        code = ord(ch)
        if ch == "\ufffd":
            replacement_count += 1
        if 0xE000 <= code <= 0xF8FF:
            private_use_count += 1
        if ch not in ("\n", "\r", "\t") and unicodedata.category(ch) == "Cc":
            control_count += 1
        if _is_expected_char(ch):
            expected_count += 1

    non_expected_count = max(0, total_chars - expected_count)
    control_ratio = (control_count / total_chars) if total_chars > 0 else 0.0
    private_ratio = (private_use_count / total_chars) if total_chars > 0 else 0.0
    non_expected_ratio = (non_expected_count / total_chars) if total_chars > 0 else 0.0

    table_applicable = source_ext == ".pdf"
    if table_applicable:
        table_stats = _table_stats(text)
        table_ratio = table_stats["table_ratio"]
        table_blocks = table_stats["table_blocks"]
    else:
        table_ratio = 0.0
        table_blocks = 0

    weird_replacement = replacement_count >= WEIRD_REPLACEMENT_THRESHOLD
    weird_control = control_ratio >= WEIRD_CONTROL_RATIO
    weird_private = private_ratio >= WEIRD_PRIVATE_RATIO
    weird_nonexpected = non_expected_ratio >= WEIRD_NONEXPECTED_RATIO
    is_weird = any([weird_replacement, weird_control, weird_private, weird_nonexpected])

    if table_applicable:
        is_table_heavy = (
            table_ratio >= TABLE_RATIO_THRESHOLD
            or (table_ratio >= TABLE_RATIO_SECONDARY and table_blocks >= TABLE_COUNT_THRESHOLD)
        )
    else:
        is_table_heavy = False

    structure_applicable = source_ext == ".pdf"
    if structure_applicable:
        empty_enum_run = _max_empty_circled_enum_run(text)
        is_structure_suspicious = empty_enum_run >= EMPTY_ENUM_RUN_MIN
    else:
        empty_enum_run = 0
        is_structure_suspicious = False

    flag_reasons = []
    if is_weird:
        flag_reasons.append("invalid_characters")
    if table_applicable and is_table_heavy:
        flag_reasons.append("high_no_tables")
    if structure_applicable and is_structure_suspicious:
        flag_reasons.append("format_mismatch_review_needed")

    return {
        "total_chars": total_chars,
        "replacement_count": replacement_count,
        "control_count": control_count,
        "private_use_count": private_use_count,
        "non_expected_ratio": round(non_expected_ratio, 6),
        "table_blocks": table_blocks,
        "table_ratio": round(table_ratio, 6),
        "empty_enum_run": empty_enum_run,
        "is_structure_suspicious": is_structure_suspicious,
        "is_weird": is_weird,
        "is_table_heavy": is_table_heavy,
        "table_applicable": table_applicable,
        "structure_applicable": structure_applicable,
        "flag_reason": "|".join(flag_reasons),
    }


def init_log():
    """Initialize CSV log file in the results folder."""
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    with open(LOG_FILE, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["timestamp", "file", "pages", "status", "error"])

def init_quality_log():
    """Initialize quality CSV log files (JP + EN) in the results folder."""
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    with open(QUALITY_LOG_FILE, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow([
            "タイムスタンプ",
            "ファイル",
            "ページ数",
            "処理結果",
            "エラー",
            "文字数",
            "置換文字(�)数",
            "制御文字数",
            "私用領域文字数",
            "想定外文字比率",
            "表ブロック数",
            "表比率",
            "空番号連続数",
            "文字化け疑い",
            "表中心",
            "構造崩れ疑い",
            "問題の種類",
            "要確認フラグ",
        ])

    with open(QUALITY_LOG_FILE_EN, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow([
            "timestamp",
            "file",
            "pages",
            "status",
            "error",
            "total_chars",
            "replacement_char_count",
            "control_char_count",
            "private_use_char_count",
            "non_expected_ratio",
            "table_blocks",
            "table_ratio",
            "empty_enum_run_max",
            "weird_text_suspected",
            "table_heavy",
            "structure_suspicious",
            "issue_types",
            "needs_review",
        ])

    _write_quality_legend_files(RESULTS_DIR)


def append_log(file_path: Path, pages: int, status: str, error: str = ""):
    """Append a single record to the log file."""
    with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            datetime.now().isoformat(timespec="seconds"),
            str(file_path),
            pages,
            status,
            error
        ])

def _jp_bool(value: bool) -> str:
    return "はい" if value else "いいえ"


def _jp_status(status: str) -> str:
    return "成功" if status == "OK" else "失敗"


def _jp_flag_reason(reason: str) -> str:
    if not reason:
        return ""
    mapping = {
        "format_mismatch_review_needed": "形式不一致",
        "invalid_characters": "無効文字",
        "high_no_tables": "表が多い",
    }
    parts = [mapping.get(r, r) for r in reason.split("|") if r]
    return " / ".join(parts)


def _jp_bool_or_na(value: bool, applicable: bool) -> str:
    if not applicable:
        return "対象外"
    return _jp_bool(value)


def append_quality_log(file_path: Path, pages: int, status: str, error: str, quality=None):
    """Append a single quality record to the quality log file."""
    if quality is None:
        quality = {
            "total_chars": 0,
            "replacement_count": 0,
            "control_count": 0,
            "private_use_count": 0,
            "non_expected_ratio": 0.0,
            "table_blocks": 0,
            "table_ratio": 0.0,
            "empty_enum_run": 0,
            "is_structure_suspicious": False,
            "is_weird": False,
            "is_table_heavy": False,
            "table_applicable": True,
            "structure_applicable": True,
            "flag_reason": "",
        }

    table_applicable = quality.get("table_applicable", True)
    structure_applicable = quality.get("structure_applicable", True)
    table_blocks = quality["table_blocks"] if table_applicable else ""
    table_ratio = quality["table_ratio"] if table_applicable else ""
    empty_enum_run = quality["empty_enum_run"] if structure_applicable else ""
    needs_review = "☆" if (status != "OK" or quality["flag_reason"]) else ""

    issue_type_jp = "抽出失敗" if status != "OK" else _jp_flag_reason(quality["flag_reason"])

    with open(QUALITY_LOG_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            datetime.now().isoformat(timespec="seconds"),
            str(file_path),
            pages,
            _jp_status(status),
            error,
            quality["total_chars"],
            quality["replacement_count"],
            quality["control_count"],
            quality["private_use_count"],
            quality["non_expected_ratio"],
            table_blocks,
            table_ratio,
            empty_enum_run,
            _jp_bool(quality["is_weird"]),
            _jp_bool_or_na(quality["is_table_heavy"], table_applicable),
            _jp_bool_or_na(quality["is_structure_suspicious"], structure_applicable),
            issue_type_jp,
            needs_review,
        ])


def _write_quality_legend_files(base_dir: Path):
    jp_path = base_dir / "extraction_quality_legend_ja.txt"
    en_path = base_dir / "extraction_quality_legend_en.txt"

    jp_text = """品質CSVの説明 (extraction_quality.csv)

- 文字数: 抽出テキスト内の文字数
- 置換文字(�)数: 文字化けの強い兆候となる置換文字の数
- 制御文字数: 改行/タブ以外の制御文字数
- 私用領域文字数: Unicode私用領域(U+E000〜U+F8FF)の文字数
- 想定外文字比率: 期待される文字以外の割合
- 表ブロック数/表比率: PDFのみ対象。表と判定された領域の数/割合
- 空番号連続数: ②③④…のような“空の番号行”が連続する最大数
- 文字化け疑い: 文字化けの兆候がある場合「はい」
- 表中心: PDFのみ対象。表が多すぎる場合「はい」、PDF以外は「対象外」
- 構造崩れ疑い: 空番号の連続が一定以上の場合「はい」(PDFのみ対象)
- 問題の種類: 次の3分類のみ
  - 形式不一致
  - 無効文字
  - 表が多い
- 抽出失敗時は「抽出失敗」と表示
- 要確認フラグ: 問題がある場合「☆」(抽出失敗も含む)

比率は「該当文字数 ÷ 文字数」で計算しています。
"""

    en_text = """Quality CSV legend (extraction_quality_en.csv)

- total_chars: total character count in extracted text
- replacement_char_count: number of replacement chars (�)
- control_char_count: control characters excluding \\n/\\r/\\t
- private_use_char_count: Unicode private-use chars (U+E000–U+F8FF)
- non_expected_ratio: ratio of characters outside the expected set
- table_blocks/table_ratio: PDF only. Count/ratio of table regions
- empty_enum_run_max: max consecutive empty circled-number lines (e.g., ②③④…)
- weird_text_suspected: "Yes" if text looks garbled
- table_heavy: "Yes" if tables dominate; "N/A" for non-PDF
- structure_suspicious: "Yes" if empty circled-number runs exceed threshold (PDF only)
- issue_types: only these three labels
  - Format Mismatch
  - Invalid Characters
  - High number of tables
- If extraction fails, issue_types shows "Extraction Failed"
- needs_review: "☆" if any issue exists (including extraction failure)

Ratios are calculated as (matching chars) / (total chars).
"""

    jp_path.write_text(jp_text, encoding="utf-8-sig")
    en_path.write_text(en_text, encoding="utf-8-sig")


def _en_bool(value: bool) -> str:
    return "Yes" if value else "No"


def _en_status(status: str) -> str:
    return "Success" if status == "OK" else "Fail"


def _en_flag_reason(reason: str) -> str:
    if not reason:
        return ""
    mapping = {
        "format_mismatch_review_needed": "Format Mismatch",
        "invalid_characters": "Invalid Characters",
        "high_no_tables": "High number of tables",
    }
    parts = [mapping.get(r, r) for r in reason.split("|") if r]
    return " / ".join(parts)


def _en_bool_or_na(value: bool, applicable: bool) -> str:
    if not applicable:
        return "N/A"
    return _en_bool(value)


def append_quality_log_en(file_path: Path, pages: int, status: str, error: str, quality=None):
    """Append a single quality record to the English quality log file."""
    if quality is None:
        quality = {
            "total_chars": 0,
            "replacement_count": 0,
            "control_count": 0,
            "private_use_count": 0,
            "non_expected_ratio": 0.0,
            "table_blocks": 0,
            "table_ratio": 0.0,
            "empty_enum_run": 0,
            "is_structure_suspicious": False,
            "is_weird": False,
            "is_table_heavy": False,
            "table_applicable": True,
            "structure_applicable": True,
            "flag_reason": "",
        }

    table_applicable = quality.get("table_applicable", True)
    structure_applicable = quality.get("structure_applicable", True)
    table_blocks = quality["table_blocks"] if table_applicable else ""
    table_ratio = quality["table_ratio"] if table_applicable else ""
    empty_enum_run = quality["empty_enum_run"] if structure_applicable else ""
    needs_review = "☆" if (status != "OK" or quality["flag_reason"]) else ""

    issue_type_en = "Extraction Failed" if status != "OK" else _en_flag_reason(quality["flag_reason"])

    with open(QUALITY_LOG_FILE_EN, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            datetime.now().isoformat(timespec="seconds"),
            str(file_path),
            pages,
            _en_status(status),
            error,
            quality["total_chars"],
            quality["replacement_count"],
            quality["control_count"],
            quality["private_use_count"],
            quality["non_expected_ratio"],
            table_blocks,
            table_ratio,
            empty_enum_run,
            _en_bool(quality["is_weird"]),
            _en_bool_or_na(quality["is_table_heavy"], table_applicable),
            _en_bool_or_na(quality["is_structure_suspicious"], structure_applicable),
            issue_type_en,
            needs_review,
        ])

def collect_supported_files(input_dir: Path):
    """Return a sorted list of all supported files under input_dir."""
    return sorted(
        (
            path
            for path in input_dir.rglob("*")
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ),
        key=lambda p: str(p).lower(),
    )


def process_documents(input_dir: Path, output_dir: Path):
    files = collect_supported_files(input_dir)
    if not files:
        print(f"No supported files (.pdf, .docx, .xls, .xlsx) found in: {input_dir}")
        return

    init_log()
    init_quality_log()
    print(f"Found {len(files)} files (PDF/Word/Excel). Starting extraction...\n")

    output_dir.mkdir(parents=True, exist_ok=True)

    with ProcessPoolExecutor() as executor:
        futures = {
            executor.submit(process_one_file, path, input_dir, output_dir): path
            for path in files
        }

        for fut in as_completed(futures):
            rel, pages, status, error, quality = fut.result()
            append_log(rel, pages, status, error)
            append_quality_log(rel, pages, status, error, quality)
            append_quality_log_en(rel, pages, status, error, quality)
            if status == "OK":
                print(f"✅ Extracted: {rel}  ({pages} pages/sheets)")
            else:
                print(f"❌ Failed: {rel} ({error})")

    print("\nAll done!")
    print(f"📂 Extracted texts: {output_dir}")
    print(f"📄 Log file: {LOG_FILE}")


# Backwards-compatible entry point name
def process_pdfs(input_dir: Path, output_dir: Path):
    process_documents(input_dir, output_dir)


if __name__ == "__main__":
    process_documents(INPUT_DIR, OUTPUT_DIR)
